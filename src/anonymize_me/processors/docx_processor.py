"""DOCX anonymization utilities."""

from __future__ import annotations

import io

from docx import Document

from .types import AnonymizedAttachment
from ..anonymizer import TextProvider


class DocxProcessor:
    """Anonymize Microsoft Word ``.docx`` files."""

    def __init__(self, anonymize: TextProvider) -> None:
        self._anonymize = anonymize

    def anonymize(self, filename: str, payload: bytes) -> AnonymizedAttachment:
        document = Document(io.BytesIO(payload))

        for paragraph in document.paragraphs:
            paragraph.text = self._anonymize(paragraph.text, context="Word paragraph")

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell.text = self._anonymize(cell.text, context="Word table cell")

        buffer = io.BytesIO()
        document.save(buffer)
        return AnonymizedAttachment(
            filename=filename,
            content=buffer.getvalue(),
            maintype="application",
            subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
