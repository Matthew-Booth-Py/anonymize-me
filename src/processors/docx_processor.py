"""DOCX anonymization utilities."""

from __future__ import annotations

import io
import uuid

from docx import Document
from docx.document import Document as DocumentType
from docx.table import Table
from docx.text.paragraph import Paragraph

from .types import AnonymizedAttachment
from ..anonymizer import ReplacementProvider, apply_replacements


class DocxProcessor:
    """Anonymize Microsoft Word ``.docx`` files using Presidio."""

    def __init__(self, replacement_provider: ReplacementProvider) -> None:
        self._replacement_provider = replacement_provider

    def anonymize(self, filename: str, payload: bytes) -> AnonymizedAttachment:
        """Anonymize DOCX content while preserving all formatting and metadata."""
        document = Document(io.BytesIO(payload))

        # Extract all text from the document to analyze with Presidio
        full_text = self._extract_all_text(document)

        # Use Presidio to generate replacements
        replacements = self._replacement_provider(full_text, context="DOCX attachment")

        # Single pass through all document elements
        self._apply_replacements_to_document(document, replacements)

        # Save with all metadata preserved
        buffer = io.BytesIO()
        document.save(buffer)

        # Generate random filename
        random_filename = f"{uuid.uuid4().hex[:12]}.docx"

        return AnonymizedAttachment(
            filename=random_filename,
            content=buffer.getvalue(),
            maintype="application",
            subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    def _extract_all_text(self, document: DocumentType) -> str:
        """Extract all text from the document for Presidio analysis."""
        text_parts = []

        # Extract from paragraphs (main body)
        for paragraph in document.paragraphs:
            if paragraph.text:
                text_parts.append(paragraph.text)

        # Extract from tables
        for table in document.tables:
            text_parts.extend(self._extract_text_from_table(table))

        # Extract from headers/footers
        for section in document.sections:
            for paragraph in section.header.paragraphs:
                if paragraph.text:
                    text_parts.append(paragraph.text)
            for paragraph in section.footer.paragraphs:
                if paragraph.text:
                    text_parts.append(paragraph.text)

        return "\n".join(text_parts)

    def _extract_text_from_table(self, table: Table) -> list[str]:
        """Extract text from a table and any nested tables."""
        text_parts = []
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text:
                        text_parts.append(paragraph.text)
                # Handle nested tables
                for nested_table in cell.tables:
                    text_parts.extend(self._extract_text_from_table(nested_table))
        return text_parts

    def _apply_replacements_to_document(
        self, document: DocumentType, replacements: dict[str, str]
    ) -> None:
        """Apply replacements to all text elements in the document in a single pass."""
        if not replacements:
            return

        # Process all paragraphs (main body)
        for paragraph in document.paragraphs:
            self._apply_replacements_to_paragraph(paragraph, replacements)

        # Process all tables
        for table in document.tables:
            self._apply_replacements_to_table(table, replacements)

        # Process all sections (headers/footers)
        for section in document.sections:
            # Header paragraphs
            for paragraph in section.header.paragraphs:
                self._apply_replacements_to_paragraph(paragraph, replacements)

            # Footer paragraphs
            for paragraph in section.footer.paragraphs:
                self._apply_replacements_to_paragraph(paragraph, replacements)

    def _apply_replacements_to_paragraph(
        self, paragraph: Paragraph, replacements: dict[str, str]
    ) -> None:
        """Apply replacements to all runs in a paragraph."""
        for run in paragraph.runs:
            if run.text:
                run.text = apply_replacements(run.text, replacements)

    def _apply_replacements_to_table(
        self, table: Table, replacements: dict[str, str]
    ) -> None:
        """Apply replacements to all cells in a table."""
        for row in table.rows:
            for cell in row.cells:
                # Process all paragraphs in the cell
                for paragraph in cell.paragraphs:
                    self._apply_replacements_to_paragraph(paragraph, replacements)

                # Process nested tables (tables within cells)
                for nested_table in cell.tables:
                    self._apply_replacements_to_table(nested_table, replacements)
